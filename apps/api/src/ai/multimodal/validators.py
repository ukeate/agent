"""
内容验证器
"""

import mimetypes
from pathlib import Path
from typing import Optional, Tuple
import hashlib
from .types import ContentType
from .config import ModelConfig

class ContentValidator:
    """内容验证器"""
    
    @staticmethod
    def validate_file(
        file_path: str,
        expected_type: Optional[ContentType] = None
    ) -> Tuple[bool, Optional[str], Optional[ContentType]]:
        """
        验证文件
        返回: (是否有效, 错误消息, 内容类型)
        """
        file_path = Path(file_path)
        
        # 检查文件是否存在
        if not file_path.exists():
            return False, f"文件不存在: {file_path}", None
        
        # 检查文件大小
        file_size = file_path.stat().st_size
        if file_size == 0:
            return False, "文件大小为0", None
        
        if file_size > ModelConfig.MAX_FILE_SIZE:
            return False, f"文件大小超过限制: {file_size} > {ModelConfig.MAX_FILE_SIZE}", None
        
        # 检测文件类型
        detected_type = ContentValidator.detect_content_type(file_path)
        
        if not detected_type:
            return False, f"无法识别的文件类型: {file_path.suffix}", None
        
        # 如果指定了期望类型，进行验证
        if expected_type and detected_type != expected_type:
            return False, f"文件类型不匹配: 期望 {expected_type}, 实际 {detected_type}", None
        
        # 检查文件格式是否支持
        if not ContentValidator.is_format_supported(file_path.suffix, detected_type):
            return False, f"不支持的文件格式: {file_path.suffix}", None
        
        return True, None, detected_type
    
    @staticmethod
    def detect_content_type(file_path: Path) -> Optional[ContentType]:
        """检测内容类型"""
        mime_type, _ = mimetypes.guess_type(str(file_path))
        suffix = file_path.suffix.lower()
        
        # 基于MIME类型检测
        if mime_type:
            if mime_type.startswith('image/'):
                return ContentType.IMAGE
            elif mime_type.startswith('video/'):
                return ContentType.VIDEO
            elif mime_type.startswith('audio/'):
                return ContentType.AUDIO
            elif mime_type == 'application/pdf':
                return ContentType.DOCUMENT
            elif mime_type.startswith('text/'):
                return ContentType.DOCUMENT
        
        # 基于文件扩展名检测
        if suffix in ModelConfig.SUPPORTED_IMAGE_FORMATS:
            return ContentType.IMAGE
        elif suffix in ModelConfig.SUPPORTED_DOCUMENT_FORMATS:
            return ContentType.DOCUMENT
        elif suffix in ModelConfig.SUPPORTED_VIDEO_FORMATS:
            return ContentType.VIDEO
        elif suffix in ModelConfig.SUPPORTED_AUDIO_FORMATS:
            return ContentType.AUDIO
        elif suffix in {'.txt', '.md', '.csv', '.json', '.xml', '.yaml', '.yml'}:
            return ContentType.TEXT
        
        return None
    
    @staticmethod
    def is_format_supported(file_extension: str, content_type: ContentType) -> bool:
        """检查文件格式是否支持"""
        return ModelConfig.is_format_supported(file_extension, content_type.value)
    
    @staticmethod
    def validate_image(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证图像文件"""
        try:
            from PIL import Image
            
            # 尝试打开图像
            with Image.open(file_path) as img:
                # 检查图像尺寸
                width, height = img.size
                if width * height > 100_000_000:  # 100MP限制
                    return False, "图像分辨率过高"
                
                # 检查图像模式
                if img.mode not in ['RGB', 'RGBA', 'L', 'P']:
                    return False, f"不支持的图像模式: {img.mode}"
            
            return True, None
            
        except Exception as e:
            return False, f"无效的图像文件: {str(e)}"
    
    @staticmethod
    def validate_document(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证文档文件"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return ContentValidator.validate_pdf(str(file_path))
        elif suffix == '.docx':
            return ContentValidator.validate_docx(str(file_path))
        elif suffix in {'.txt', '.md', '.csv'}:
            return ContentValidator.validate_text(str(file_path))
        else:
            return False, f"不支持的文档格式: {suffix}"
    
    @staticmethod
    def validate_pdf(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证PDF文件"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # 检查是否加密
                if pdf_reader.is_encrypted:
                    return False, "PDF文件已加密"
                
                # 检查页数
                page_count = len(pdf_reader.pages)
                if page_count == 0:
                    return False, "PDF文件没有页面"
                if page_count > 1000:
                    return False, f"PDF页数过多: {page_count}"
            
            return True, None
            
        except Exception as e:
            return False, f"无效的PDF文件: {str(e)}"
    
    @staticmethod
    def validate_docx(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证DOCX文件"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 检查是否有内容
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                return False, "DOCX文件没有内容"
            
            return True, None
            
        except Exception as e:
            return False, f"无效的DOCX文件: {str(e)}"
    
    @staticmethod
    def validate_text(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                # 尝试读取前1KB
                content = f.read(1024)
                
                # 检查是否为空
                if not content.strip():
                    return False, "文本文件为空"
            
            return True, None
            
        except UnicodeDecodeError:
            return False, "文件编码不是UTF-8"
        except Exception as e:
            return False, f"无法读取文本文件: {str(e)}"
    
    @staticmethod
    def validate_video(file_path: str) -> Tuple[bool, Optional[str]]:
        """验证视频文件"""
        try:
            import cv2
            
            cap = cv2.VideoCapture(file_path)
            
            if not cap.isOpened():
                return False, "无法打开视频文件"
            
            # 检查视频属性
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            fps = cap.get(cv2.CAP_PROP_FPS)
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            cap.release()
            
            if frame_count == 0:
                return False, "视频没有帧"
            
            if width * height > 8294400:  # 4K限制
                return False, f"视频分辨率过高: {width}x{height}"
            
            duration = frame_count / fps if fps > 0 else 0
            if duration > 3600:  # 1小时限制
                return False, f"视频时长过长: {duration}秒"
            
            return True, None
            
        except Exception as e:
            return False, f"无效的视频文件: {str(e)}"
    
    @staticmethod
    def calculate_file_hash(file_path: str) -> str:
        """计算文件哈希"""
        hasher = hashlib.sha256()
        
        with open(file_path, 'rb') as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        
        return hasher.hexdigest()
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """清理文件名"""
        # 移除不安全字符
        safe_chars = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789.-_')
        sanitized = ''.join(c if c in safe_chars else '_' for c in filename)
        
        # 限制长度
        if len(sanitized) > 255:
            name, ext = sanitized.rsplit('.', 1) if '.' in sanitized else (sanitized, '')
            sanitized = name[:250] + '.' + ext if ext else name[:255]
        
        return sanitized
